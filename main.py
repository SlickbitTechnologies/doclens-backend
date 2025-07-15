import os
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
import google.generativeai as genai
import re
import json
import difflib

def fuzzy_match(a, b):
    if not a or not b:
        return False
    a, b = a.lower(), b.lower()
    return a == b or a in b or b in a

def normalize_title(title):
    # Remove numbers, punctuation, and convert to uppercase
    return re.sub(r'[^A-Z ]', '', title.upper()) if title else ''

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://delightful-sky-044169710.2.azurestaticapps.net",  # Azure frontend
        "http://localhost:5173",  # Local development (optional)
        "http://127.0.0.1:5173"   # Local development (optional)
        ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/compare")
async def compare(source: UploadFile = File(...), child: UploadFile = File(...)):
    try:
        # Read file contents
        source.file.seek(0)
        cds_text = source.file.read().decode(errors='ignore') if source.filename.endswith('.txt') else source.file.read().decode(errors='ignore')
        child.file.seek(0)
        child_text = child.file.read().decode(errors='ignore') if child.filename.endswith('.txt') else child.file.read().decode(errors='ignore')

        # If PDF or DOCX, use PyPDF2 or python-docx to extract text
        if source.filename.endswith('.pdf'):
            import PyPDF2
            source.file.seek(0)
            reader = PyPDF2.PdfReader(source.file)
            cds_text = "".join([page.extract_text() or "" for page in reader.pages])
        elif source.filename.endswith('.docx'):
            from docx import Document
            source.file.seek(0)
            doc = Document(source.file)
            cds_text = "\n".join([para.text for para in doc.paragraphs])

        if child.filename.endswith('.pdf'):
            import PyPDF2
            child.file.seek(0)
            reader = PyPDF2.PdfReader(child.file)
            child_text = "".join([page.extract_text() or "" for page in reader.pages])
        elif child.filename.endswith('.docx'):
            from docx import Document
            child.file.seek(0)
            doc = Document(child.file)
            child_text = "\n".join([para.text for para in doc.paragraphs])

        # Gemini API key
        gemini_api_key = os.environ.get("GEMINI_API_KEY")
        if not gemini_api_key:
            raise HTTPException(status_code=500, detail="Gemini API key is not set.")
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')

        # Single Gemini call for extraction, matching, and summarization
        comparison_prompt = f"""
You are an expert in regulatory document analysis. Here are two documents: a Core Data Sheet (CDS) and a Child (USPI, SmPC, etc.) document.

Instructions:
1. Split each document into sections and sub-sections based on ALL headings and subheadings (e.g., 1, 1.1, 1.2, etc.).
2. For each section, extract and return the FULL, unaltered content from the document. Do NOT summarize, condense, or omit any part of the section content. Return the complete text as it appears in the document.
3. Ensure that section names (titles) are not repeated in the output; each section should appear only once per document.
4. For each section in the child document, find the best-matching section in the CDS (by meaning, not just wording).
5. For each matched pair, return:
   - cds_title, cds_content
   - child_title, child_content
   - similarity_score (0-1)
   - summary: a 1-2 sentence summary of the key differences
   - missing_in_child: list of sentences present in CDS but missing in Child
   - missing_in_cds: list of sentences present in Child but missing in CDS
6. For unmatched sections in either document, include them with empty fields for the missing side.

Respond ONLY with a JSON array of objects as described above. No explanation, no markdown, no extra text, no code block markers.

CDS Document:
{cds_text}

Child Document:
{child_text}
"""
        comparison_response = model.generate_content(comparison_prompt)
        match = re.search(r'\[.*\]', comparison_response.text, re.DOTALL)
        comparison_json = match.group(0) if match else comparison_response.text
        try:
            matched_pairs = json.loads(comparison_json)
        except Exception as e:
            # print("Failed to parse Gemini comparison response as JSON:", e)
            matched_pairs = []

        # Build section lists for UI compatibility
        cds_sections = []
        child_sections = []
        for pair in matched_pairs:
            if pair.get('cds_title') and pair.get('cds_content'):
                cds_sections.append({'title': pair['cds_title'], 'content': pair['cds_content']})
            if pair.get('child_title') and pair.get('child_content'):
                child_sections.append({'title': pair['child_title'], 'content': pair['child_content']})

        # Build unified list for all sections (matched and unmatched)
        unified_list = []
        for pair in matched_pairs:
            unified_list.append({
                'cds_title': pair.get('cds_title', ''),
                'cds_content': pair.get('cds_content', ''),
                'child_title': pair.get('child_title', ''),
                'child_content': pair.get('child_content', ''),
                'similarity': pair.get('similarity_score'),
                'status': 'Matched' if pair.get('cds_title') and pair.get('child_title') else 'Unmatched'
            })

        # Build section_comparisons for overall/section-wise summary
        section_comparisons = []
        for pair in matched_pairs:
            section_comparisons.append({
                'title': pair.get('cds_title') or pair.get('child_title'),
                'change_count': len(pair.get('missing_in_child') or []) + len(pair.get('missing_in_cds') or []),
                'summary': pair.get('summary', ''),
                'diff': '',  # Optionally, you can generate a diff here if needed
                'missing_in_child': pair.get('missing_in_child', []),
                'missing_in_cds': pair.get('missing_in_cds', [])
            })

        return {
            "matched_pairs": matched_pairs,
            "cds_sections": cds_sections,
            "child_sections": child_sections,
            "unified_list": unified_list,
            "section_comparisons": section_comparisons
        }
    except Exception as e:
        # print(f"Error in /compare: {e}")
        raise HTTPException(status_code=500, detail= "Failed to upload. Please try again later")
